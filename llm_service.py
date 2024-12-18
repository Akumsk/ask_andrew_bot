# llm_service.py

import os
import pandas as pd
from io import StringIO
from docx import Document as DocxDocument
import asyncio
from langchain.chains.history_aware_retriever import create_history_aware_retriever
from langchain.chains.retrieval import create_retrieval_chain
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder, PromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyMuPDFLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.schema import Document
from langchain.chains.combine_documents import create_stuff_documents_chain
import tiktoken

from settings import OPENAI_API_KEY, MODEL_NAME, CHAT_HISTORY_LEVEL, DOCS_IN_RETRIEVER
from helpers import current_timestamp


class LLMService:
    vector_store = None  # Class variable to store the vector store
    def __init__(self, model_name=MODEL_NAME):
        self.llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, model_name=model_name)
#        self.vector_store = None

    def load_excel_file(self, file_path):
        data = pd.read_excel(file_path)
        text_data = StringIO()
        data.to_string(buf=text_data)
        return text_data.getvalue()

    def load_word_file(self, file_path):
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    def build_context(documents, llm):
        context = ""
        # Define the prompt template
        prompt_template = (
            "Identify the type of document, what tasks this document can be used for, "
            "do not include specific details.\n\n"
            "Document Content:\n{document_content}\n\n"
            "Summary:"
        )

        # Initialize the prompt
        template = PromptTemplate(
            input_variables=["document_content"],
            template=prompt_template
        )

        for doc in documents:
            # Prepare the prompt for this document
            prompt = template.format(document_content=doc.page_content[:1000])  # Limit content to first 1000 chars

            # Get the description from the LLM
            description = llm(prompt).strip()

            # Store the description in metadata
            doc.metadata['generated_description'] = description

            # Build the context by combining descriptions and content
            context += f"Description: {description}\nContent: {doc.page_content}\n\n"

        return context

    def load_and_index_documents(self, folder_path):
        documents = []
        found_valid_file = False

        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)

            if filename.endswith(".pdf"):
                loader = PyMuPDFLoader(file_path)
                docs = loader.load()
                for doc in docs:
                    doc.metadata = {"source": filename}
                    documents.append(doc)
                found_valid_file = True

            elif filename.endswith(".docx"):
                content = self.load_word_file(file_path)
                doc = Document(page_content=content, metadata={"source": filename})
                documents.append(doc)
                found_valid_file = True

            elif filename.endswith(".xlsx"):
                content = self.load_excel_file(file_path)
                doc = Document(page_content=content, metadata={"source": filename})
                documents.append(doc)
                found_valid_file = True

        if not found_valid_file:
            return "No valid files found in the folder. Please provide PDF, Word, or Excel files."

        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        split_docs = text_splitter.split_documents(documents)

        embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
        LLMService.vector_store = FAISS.from_documents(split_docs, embeddings)
        return "Documents successfully indexed."

    def generate_response(self, prompt, chat_history=None):

        if not LLMService.vector_store:
            return (
                "Please set the folder path using /folder and ensure documents are loaded.",
                None,
            )

        # Ensure chat_history is a list
        if chat_history is None:
            chat_history = []

        # Get the top k relevant documents
#        similar_docs = get_relevant_documents(prompt, k=2)

#        if not similar_docs:
#            return "No relevant documents found.", None

        # Create the retriever
        retriever = LLMService.vector_store.as_retriever(search_kwargs={'k': DOCS_IN_RETRIEVER})

        # Create the history-aware retriever
        retriever_prompt = ChatPromptTemplate.from_messages(
            [
                MessagesPlaceholder(variable_name="chat_history"),
                ("user", "{input}"),
                (
                    "user",
                    "Given the above conversation, generate a search query to look up in order to get information relevant to the conversation",
                ),
            ]
        )

        history_aware_retriever = create_history_aware_retriever(
            llm=self.llm, retriever=retriever, prompt=retriever_prompt
        )

        # Create the question-answering chain
        system_prompt = (
            "You are a project assistant on design and construction projects. "
            "Use the following pieces of retrieved context to answer "
            "the question. If you don't know the answer, say that you "
            f"don't know. If you need to use current date, today is {current_timestamp()}."
            " Do not include references to the source documents in your answer."
            "If Prompt include request to provide a link to documents in context, respond have to be: Please follow the link below:"
            " \n\n{context}"
        )

        prompt_template = ChatPromptTemplate.from_messages(
            [
                ("system", system_prompt),
                MessagesPlaceholder(variable_name="chat_history"),
                ("user", "{input}"),
            ]
        )

        question_answer_chain = create_stuff_documents_chain(self.llm, prompt_template)

        # Create the retrieval chain using create_retrieval_chain
        rag_chain = create_retrieval_chain(
            retriever=history_aware_retriever, combine_docs_chain=question_answer_chain
        )

        # Run the chain with the provided prompt and chat history
        result = rag_chain.invoke({"input": prompt, "chat_history": chat_history})

        answer = result.get("answer", "")
        sources = result.get("context", [])

        if not sources:
            return answer, None

        source_files = set(
            [doc.metadata["source"] for doc in sources if "source" in doc.metadata]
        )

        return answer, source_files

    def count_tokens_in_context(self, folder_path):
        """Counts the total number of tokens in documents within a folder."""
        documents = []
        found_valid_file = False

        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)

            if filename.endswith(".pdf"):
                loader = PyMuPDFLoader(file_path)
                docs = loader.load()
                for doc in docs:
                    doc.metadata = {"source": filename}
                    documents.append(doc)
                found_valid_file = True

            elif filename.endswith(".docx"):
                content = self.load_word_file(file_path)
                doc = Document(page_content=content, metadata={"source": filename})
                documents.append(doc)
                found_valid_file = True

            elif filename.endswith(".xlsx"):
                content = self.load_excel_file(file_path)
                doc = Document(page_content=content, metadata={"source": filename})
                documents.append(doc)
                found_valid_file = True

        if not found_valid_file:
            return "No valid files found in the folder. Please provide PDF, Word, or Excel files."

        # Count tokens in the documents
        tokenizer = tiktoken.encoding_for_model(
            "gpt-4"
        )  # Tokenizer for the specific model

        total_tokens = 0
        for doc in documents:
            total_tokens += len(tokenizer.encode(doc.page_content))

        return total_tokens

def get_relevant_documents(query, k):
    if not LLMService.vector_store:
        return []

    similar_docs = LLMService.vector_store.similarity_search(query, k=k)
    return similar_docs